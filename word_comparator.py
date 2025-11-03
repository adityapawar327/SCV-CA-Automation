import streamlit as st
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
import io
from difflib import SequenceMatcher
import tempfile
import os
from PIL import Image, ImageDraw
import fitz  # PyMuPDF
import numpy as np
import easyocr
import cv2

# Initialize EasyOCR reader (cached)
@st.cache_resource
def get_ocr_reader():
    """Initialize and cache EasyOCR reader"""
    return easyocr.Reader(['en'], gpu=False)

def docx_to_pdf_to_image(docx_file):
    """Convert DOCX to image via PDF using LibreOffice or direct rendering"""
    # Save docx to temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
        tmp_docx.write(docx_file.read())
        docx_path = tmp_docx.name
    
    try:
        # Try to convert using python-docx and render
        doc = Document(docx_path)
        
        # Create a simple text representation and render as image
        from docx2pdf import convert
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        try:
            convert(docx_path, pdf_path)
            
            # Convert PDF to image
            pdf_doc = fitz.open(pdf_path)
            images = []
            
            for page_num in range(pdf_doc.page_count):
                page = pdf_doc[page_num]
                mat = fitz.Matrix(200/72, 200/72)  # 200 DPI
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                images.append(img)
            
            pdf_doc.close()
            os.unlink(pdf_path)
            return images
        except:
            # Fallback: create simple text image
            return create_text_image_from_docx(doc)
    finally:
        os.unlink(docx_path)

def create_text_image_from_docx(doc):
    """Create a high-quality image representation of the document"""
    from PIL import ImageFont, ImageDraw
    
    # Extract all text with better formatting
    full_text = '\n'.join([para.text for para in doc.paragraphs])
    
    # Create image with higher resolution
    img_width = 1200
    font_size = 16
    line_height = 24
    padding = 40
    
    # Estimate height
    lines = full_text.split('\n')
    img_height = max(1500, len(lines) * line_height + padding * 2)
    
    # Create white background
    img = Image.new('RGB', (img_width, img_height), color='white')
    draw = ImageDraw.Draw(img)
    
    # Try to use a better font
    try:
        font = ImageFont.truetype("arial.ttf", font_size)
    except:
        try:
            font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", font_size)
        except:
            font = ImageFont.load_default()
    
    # Draw text with better wrapping
    y_position = padding
    for line in lines:
        if not line.strip():
            y_position += line_height // 2
            continue
            
        # Wrap long lines
        words = line.split()
        current_line = ""
        
        for word in words:
            test_line = current_line + word + " "
            try:
                bbox = draw.textbbox((0, 0), test_line, font=font)
                text_width = bbox[2] - bbox[0]
            except:
                text_width = len(test_line) * 8
            
            if text_width < img_width - padding * 2:
                current_line = test_line
            else:
                if current_line:
                    draw.text((padding, y_position), current_line.strip(), fill='black', font=font)
                    y_position += line_height
                current_line = word + " "
        
        if current_line:
            draw.text((padding, y_position), current_line.strip(), fill='black', font=font)
            y_position += line_height
    
    # Crop to actual content
    if y_position < img_height:
        img = img.crop((0, 0, img_width, y_position + padding))
    
    return [img]

def extract_text_with_boxes_from_image(image):
    """Extract text and bounding boxes using OCR"""
    reader = get_ocr_reader()
    
    # Convert PIL to numpy
    img_array = np.array(image)
    
    # Perform OCR
    results = reader.readtext(img_array, detail=1)
    
    text_blocks = []
    for (bbox, text, conf) in results:
        if conf < 0.3:  # Skip low confidence
            continue
            
        x_coords = [point[0] for point in bbox]
        y_coords = [point[1] for point in bbox]
        
        left = int(min(x_coords))
        top = int(min(y_coords))
        width = int(max(x_coords) - left)
        height = int(max(y_coords) - top)
        
        text_blocks.append({
            'text': text.strip(),
            'left': left,
            'top': top,
            'width': width,
            'height': height,
            'conf': conf
        })
    
    return text_blocks

def extract_text_from_docx(docx_file):
    """Extract text from Word document with paragraph structure"""
    doc = Document(docx_file)
    paragraphs = []
    
    for para in doc.paragraphs:
        # Include all paragraphs, even empty ones for accurate positioning
        paragraphs.append({
            'text': para.text,
            'style': para.style.name,
            'is_empty': not para.text.strip()
        })
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        paragraphs.append({
                            'text': para.text,
                            'style': 'Table',
                            'is_empty': False
                        })
    
    return paragraphs, doc

def normalize_text_for_comparison(text):
    """Normalize text for better comparison"""
    import re
    # Remove extra whitespace
    text = ' '.join(text.split())
    # Normalize punctuation
    text = re.sub(r'\s+([.,;:!?])', r'\1', text)
    return text.strip()

def compare_text_blocks_word_level(blocks1, blocks2):
    """Compare text blocks at word level and find differences"""
    # Extract all words with their positions
    words1 = []
    for block in blocks1:
        words_in_block = block['text'].split()
        for word in words_in_block:
            words1.append({
                'word': normalize_text_for_comparison(word),
                'original': word,
                'block': block
            })
    
    words2 = []
    for block in blocks2:
        words_in_block = block['text'].split()
        for word in words_in_block:
            words2.append({
                'word': normalize_text_for_comparison(word),
                'original': word,
                'block': block
            })
    
    # Compare word sequences
    text1 = ' '.join([w['word'] for w in words1])
    text2 = ' '.join([w['word'] for w in words2])
    
    matcher = SequenceMatcher(None, text1, text2, autojunk=False)
    opcodes = matcher.get_opcodes()
    
    diff_blocks = []
    diff_blocks_set = set()
    
    for tag, i1, i2, j1, j2 in opcodes:
        if tag in ['replace', 'insert', 'delete']:
            # Find which blocks in text2 contain differences
            char_pos = 0
            for idx, word_info in enumerate(words2):
                word_start = char_pos
                word_end = char_pos + len(word_info['word']) + 1
                
                # Check if this word overlaps with difference
                if (word_start <= j1 < word_end) or (word_start < j2 <= word_end) or (j1 <= word_start and word_end <= j2):
                    block_id = id(word_info['block'])
                    if block_id not in diff_blocks_set:
                        diff_blocks.append(word_info['block'])
                        diff_blocks_set.add(block_id)
                
                char_pos = word_end
    
    return diff_blocks

def compare_documents(doc1_paras, doc2_paras):
    """Compare two documents and find differences with improved accuracy"""
    differences = []
    diff_para_indices = []
    
    # Simple paragraph-by-paragraph comparison for maximum accuracy
    max_len = max(len(doc1_paras), len(doc2_paras))
    
    for i in range(max_len):
        text1 = doc1_paras[i]['text'] if i < len(doc1_paras) else ""
        text2 = doc2_paras[i]['text'] if i < len(doc2_paras) else ""
        
        # Normalize for comparison
        norm_text1 = normalize_text_for_comparison(text1)
        norm_text2 = normalize_text_for_comparison(text2)
        
        # Skip empty paragraphs in both
        if not norm_text1 and not norm_text2:
            continue
        
        # Check if different
        if norm_text1 != norm_text2:
            # Calculate similarity
            char_matcher = SequenceMatcher(None, norm_text1, norm_text2, autojunk=False)
            similarity = char_matcher.ratio()
            
            # Determine type
            if not text1:
                diff_type = 'insert'
            elif not text2:
                diff_type = 'delete'
            else:
                diff_type = 'replace'
            
            differences.append({
                'type': diff_type,
                'text1': text1,
                'text2': text2,
                'para_index': i,
                'similarity': similarity
            })
            
            # Only add to diff_para_indices if it's in doc2
            if i < len(doc2_paras):
                diff_para_indices.append(i)
    
    return differences, diff_para_indices

def estimate_paragraph_positions(image, paragraphs, diff_indices):
    """Estimate bounding boxes for paragraphs with differences using better heuristics"""
    img_height = image.height
    img_width = image.width
    
    # Filter out empty paragraphs for better positioning
    non_empty_paras = [p for p in paragraphs if not p.get('is_empty', False)]
    num_paragraphs = len(non_empty_paras)
    
    if num_paragraphs == 0:
        return []
    
    # Estimate paragraph height based on content
    avg_para_height = img_height / num_paragraphs
    
    boxes = []
    current_y = 20  # Start with some top margin
    
    for idx, para in enumerate(paragraphs):
        if para.get('is_empty', False):
            current_y += 10  # Small space for empty paragraphs
            continue
        
        # Calculate height based on text length
        text_length = len(para['text'])
        estimated_lines = max(1, text_length // 80)  # Assume ~80 chars per line
        para_height = estimated_lines * 20  # ~20 pixels per line
        
        # Check if this paragraph has differences
        if idx in diff_indices:
            padding = 10
            boxes.append({
                'left': padding,
                'top': max(0, int(current_y)),
                'width': img_width - 2 * padding,
                'height': int(para_height)
            })
        
        current_y += para_height + 5  # Add small gap between paragraphs
    
    return boxes

def annotate_image_with_boxes(image, boxes):
    """Draw black outline boxes on image"""
    img_copy = image.copy()
    draw = ImageDraw.Draw(img_copy)
    
    for box in boxes:
        x = box['left']
        y = box['top']
        w = box['width']
        h = box['height']
        
        # Draw black outline with slight padding
        padding = 2
        draw.rectangle(
            [x-padding, y-padding, x+w+padding, y+h+padding],
            outline=(0, 0, 0, 255),
            width=2
        )
    
    return img_copy

def highlight_differences_in_doc(doc, differences, text_content):
    """Create a new document with highlighted differences"""
    # Create new document
    new_doc = Document()
    
    # Copy styles from original
    for section in doc.sections:
        new_section = new_doc.sections[0]
        new_section.page_height = section.page_height
        new_section.page_width = section.page_width
        new_section.left_margin = section.left_margin
        new_section.right_margin = section.right_margin
    
    # Get paragraph indices with differences
    diff_para_indices = set()
    for diff in differences:
        if 'para_index' in diff:
            diff_para_indices.add(diff['para_index'])
    
    # Add paragraphs with highlighting
    for idx, para in enumerate(doc.paragraphs):
        new_para = new_doc.add_paragraph()
        new_para.style = para.style
        
        para_text = para.text
        
        # Check if this paragraph has differences
        if idx in diff_para_indices:
            # Highlight entire paragraph
            run = new_para.add_run(para_text)
            
            # Copy formatting from original
            if para.runs:
                orig_run = para.runs[0]
                run.bold = orig_run.bold
                run.italic = orig_run.italic
                run.underline = orig_run.underline
            
            # Highlight the entire paragraph
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            # Copy paragraph as-is
            run = new_para.add_run(para_text)
            if para.runs:
                orig_run = para.runs[0]
                run.bold = orig_run.bold
                run.italic = orig_run.italic
                run.underline = orig_run.underline
    
    return new_doc

def create_comparison_report(differences, doc1_name, doc2_name):
    """Create a text report of differences"""
    report = []
    report.append(f"Comparison Report: {doc1_name} vs {doc2_name}")
    report.append("=" * 80)
    report.append(f"\nTotal Differences: {len(differences)}\n")
    
    for idx, diff in enumerate(differences, 1):
        report.append(f"\nDifference #{idx} ({diff['type'].upper()}):")
        report.append("-" * 40)
        
        if diff['type'] == 'replace':
            report.append(f"Original: {diff['text1'][:100]}...")
            report.append(f"Changed:  {diff['text2'][:100]}...")
        elif diff['type'] == 'delete':
            report.append(f"Deleted:  {diff['text1'][:100]}...")
        elif diff['type'] == 'insert':
            report.append(f"Added:    {diff['text2'][:100]}...")
    
    return '\n'.join(report)

def run_word_comparator():
    st.title("📄 Word Document Comparator")
    st.caption("Compare Word documents with precise visual highlighting")
    
    # Initialize session state
    if 'word_files_data' not in st.session_state:
        st.session_state.word_files_data = {}
    if 'word_comparison_done' not in st.session_state:
        st.session_state.word_comparison_done = False
    
    st.divider()
    
    uploaded_files = st.file_uploader(
        "Upload Word documents (up to 10)",
        type=['docx'],
        accept_multiple_files=True,
        help="Upload multiple Word documents to compare them"
    )
    
    if uploaded_files:
        if len(uploaded_files) > 10:
            st.warning("Maximum 10 files allowed. Processing first 10 files.")
            uploaded_files = uploaded_files[:10]
        
        st.info(f"Loaded {len(uploaded_files)} file(s)")
        
        # Load all files
        file_names = [f.name for f in uploaded_files]
        if set(file_names) != set(st.session_state.word_files_data.keys()):
            with st.spinner("Loading documents..."):
                word_files_data = {}
                for file in uploaded_files:
                    paragraphs, doc = extract_text_from_docx(file)
                    word_files_data[file.name] = {
                        'paragraphs': paragraphs,
                        'doc': doc,
                        'file': file
                    }
                
                st.session_state.word_files_data = word_files_data
                st.session_state.word_comparison_done = False
        
        if len(st.session_state.word_files_data) < 2:
            st.error("Need at least 2 valid files to compare")
            return
        
        # File selection for comparison
        st.markdown("### Select Files to Compare")
        available_files = list(st.session_state.word_files_data.keys())
        
        col1, col2 = st.columns(2)
        with col1:
            file1 = st.selectbox("Base Document", available_files, key='word_base_file')
        with col2:
            compare_files = [f for f in available_files if f != file1]
            selected_compare_files = st.multiselect(
                "Compare Against",
                compare_files,
                default=compare_files[:1] if compare_files else [],
                key='word_compare_files'
            )
        
        if not selected_compare_files:
            st.warning("Please select at least one file to compare against")
            return
        
        if st.button("📝 Compare Documents", type="primary", use_container_width=True):
            st.session_state.word_comparison_done = True
            st.session_state.selected_word_base = file1
            st.session_state.selected_word_compare = selected_compare_files
        
        if st.session_state.word_comparison_done:
            base_file = st.session_state.selected_word_base
            compare_files = st.session_state.selected_word_compare
            
            base_data = st.session_state.word_files_data[base_file]
            
            # Create tabs for each comparison
            tabs = st.tabs(compare_files + ["Summary"])
            
            all_differences = []
            
            for idx, compare_file in enumerate(compare_files):
                with tabs[idx]:
                    st.markdown(f"### {base_file} vs {compare_file}")
                    
                    compare_data = st.session_state.word_files_data[compare_file]
                    
                    with st.spinner("Comparing documents..."):
                        # Compare documents
                        differences, diff_para_indices = compare_documents(
                            base_data['paragraphs'],
                            compare_data['paragraphs']
                        )
                        
                        # Debug info
                        st.info(f"Found {len(differences)} difference(s) in {len(diff_para_indices)} paragraph(s)")
                        
                        # Show debug details
                        with st.expander("🔍 Debug Info"):
                            st.write(f"Base document: {len(base_data['paragraphs'])} paragraphs")
                            st.write(f"Compare document: {len(compare_data['paragraphs'])} paragraphs")
                            st.write(f"Diff paragraph indices: {diff_para_indices}")
                            
                            if differences:
                                st.write("First few differences:")
                                for idx, diff in enumerate(differences[:3]):
                                    st.write(f"Diff {idx+1}: Type={diff['type']}, Para={diff['para_index']}")
                                    st.write(f"  Text1: {diff['text1'][:100]}...")
                                    st.write(f"  Text2: {diff['text2'][:100]}...")
                        
                        # Convert documents to images
                        with st.spinner("Rendering documents and performing OCR..."):
                            # Reset file pointers
                            base_data['file'].seek(0)
                            compare_data['file'].seek(0)
                            
                            base_images = create_text_image_from_docx(base_data['doc'])
                            compare_images = create_text_image_from_docx(compare_data['doc'])
                            
                            # Work with first page
                            base_img = base_images[0]
                            compare_img = compare_images[0]
                            
                            # Extract text with bounding boxes using OCR
                            base_blocks = extract_text_with_boxes_from_image(base_img)
                            compare_blocks = extract_text_with_boxes_from_image(compare_img)
                            
                            # Find differences at word level
                            diff_blocks = compare_text_blocks_word_level(base_blocks, compare_blocks)
                            
                            # Annotate compare image with precise boxes
                            annotated_compare = annotate_image_with_boxes(compare_img, diff_blocks)
                        
                        # Display images side by side
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.markdown(f"**{base_file} (Original)**")
                            st.image(base_img, width=None)
                        
                        with col2:
                            st.markdown(f"**{compare_file} (Differences Highlighted)**")
                            st.image(annotated_compare, width=None)
                            st.caption(f"Found {len(diff_blocks)} text block(s) with differences")
                        
                        # Create highlighted document for download
                        text_content = '\n'.join([p['text'] for p in compare_data['paragraphs']])
                        highlighted_doc = highlight_differences_in_doc(
                            compare_data['doc'],
                            differences,
                            text_content
                        )
                        
                        # Download highlighted document
                        st.divider()
                        st.markdown("### 💾 Download Results")
                        
                        col1, col2 = st.columns(2)
                        
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            # Save highlighted document
                            buf = io.BytesIO()
                            highlighted_doc.save(buf)
                            buf.seek(0)
                            
                            st.download_button(
                                "⬇️ Download Highlighted Document",
                                buf.getvalue(),
                                f"highlighted_{compare_file}",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"dl_doc_{compare_file}",
                                use_container_width=True
                            )
                        
                        with col2:
                            # Download annotated image
                            buf_img = io.BytesIO()
                            annotated_compare.save(buf_img, format='PNG')
                            st.download_button(
                                "⬇️ Download Annotated Image",
                                buf_img.getvalue(),
                                f"annotated_{compare_file}.png",
                                "image/png",
                                key=f"dl_img_{compare_file}",
                                use_container_width=True
                            )
                        
                        with col3:
                            # Create comparison report
                            report = create_comparison_report(differences, base_file, compare_file)
                            st.download_button(
                                "⬇️ Download Comparison Report",
                                report,
                                f"report_{base_file}_vs_{compare_file}.txt",
                                "text/plain",
                                key=f"dl_report_{compare_file}",
                                use_container_width=True
                            )
                        
                        # Show detailed differences
                        with st.expander("📋 View Detailed Differences"):
                            for diff_idx, diff in enumerate(differences[:20], 1):  # Show first 20
                                st.markdown(f"**Difference #{diff_idx} ({diff['type'].upper()})**")
                                
                                if diff['type'] == 'replace':
                                    col1, col2 = st.columns(2)
                                    with col1:
                                        st.text_area(
                                            "Original",
                                            diff['text1'],
                                            height=100,
                                            key=f"diff_orig_{compare_file}_{diff_idx}"
                                        )
                                    with col2:
                                        st.text_area(
                                            "Changed",
                                            diff['text2'],
                                            height=100,
                                            key=f"diff_new_{compare_file}_{diff_idx}"
                                        )
                                elif diff['type'] == 'delete':
                                    st.text_area(
                                        "Deleted",
                                        diff['text1'],
                                        height=100,
                                        key=f"diff_del_{compare_file}_{diff_idx}"
                                    )
                                elif diff['type'] == 'insert':
                                    st.text_area(
                                        "Added",
                                        diff['text2'],
                                        height=100,
                                        key=f"diff_add_{compare_file}_{diff_idx}"
                                    )
                                
                                st.markdown("---")
                        
                        # Store for summary
                        all_differences.append({
                            'Base File': base_file,
                            'Compare File': compare_file,
                            'Total Differences': len(differences),
                            'Replacements': sum(1 for d in differences if d['type'] == 'replace'),
                            'Deletions': sum(1 for d in differences if d['type'] == 'delete'),
                            'Insertions': sum(1 for d in differences if d['type'] == 'insert')
                        })
            
            # Summary tab
            with tabs[-1]:
                st.markdown("### Comparison Summary")
                
                if all_differences:
                    import pandas as pd
                    summary_df = pd.DataFrame(all_differences)
                    
                    st.markdown(f"**Total Comparisons: {len(all_differences)}**")
                    st.dataframe(summary_df)
                    
                    # Export summary
                    csv = summary_df.to_csv(index=False)
                    st.download_button(
                        label="📥 Download Summary (CSV)",
                        data=csv,
                        file_name="word_comparison_summary.csv",
                        mime="text/csv",
                        use_container_width=True
                    )
                    
                    # Statistics
                    st.divider()
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Total Differences", summary_df['Total Differences'].sum())
                    with col2:
                        st.metric("Total Replacements", summary_df['Replacements'].sum())
                    with col3:
                        st.metric("Total Insertions", summary_df['Insertions'].sum())
                else:
                    st.success("No differences found")
    
    else:
        st.info("👆 Upload Word documents to begin comparison")
        
        with st.expander("ℹ️ How to use"):
            st.markdown("""
            1. Upload multiple Word documents (.docx format)
            2. Select a base document and documents to compare against
            3. Click 'Compare Documents' to see differences
            4. Yellow highlighting shows changes in the document
            5. View detailed differences in the expandable section
            6. Download highlighted documents and comparison reports
            7. View summary statistics across all comparisons
            """)
